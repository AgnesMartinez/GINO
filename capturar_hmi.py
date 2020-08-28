import PySimpleGUI as sg
import sqlite3
from datetime import date, datetime
from docx import Document
import concurrent.futures
import pymysql.cursors

#conectar a base de datos local
DEFAULT_PATH = "./base_de_datos/db.sqlite3"

def db_connect(db_path=DEFAULT_PATH):
    conn_local = sqlite3.connect(db_path)
    return conn_local

conn_local = db_connect()

#conectar a base de datos maria
def maria_connect():
    try:
        conn = pymysql.connect(
            user="agnesm",
            password="torrente",
            host="192.168.1.11",
            database="hmi",
            port=3307
        )
        return conn

    except:
        pass

#variables globales a utilizar
fechahoy = date.today()
x = datetime.today()
tiempo = x.strftime("%H:%M")
insumo = ""
output_pedia = ""
output_gine = ""
catalogo = {}
mcombo = []
alista = []
rreporte = []


with conn_local:
    query = """SELECT clave,nombre FROM catalogoinsumos ORDER BY nombre"""
    lista = conn_local.execute(query)
    for item in lista:
        catalogo[item[1]] = item[0]
        mcombo.append(item[1])
    

def reporte_gine():
    """Generar reporte de consumibles ginecologia"""

    query = """SELECT ginecologia.clave,ginecologia.fecha,ginecologia.hora,ginecologia.servicio,
    catalogoinsumos.nombre, COUNT (catalogoinsumos.nombre) cantidad
    FROM ginecologia,catalogoinsumos 
    WHERE ginecologia.CURP=?
    AND ginecologia.clave = catalogoinsumos.clave AND ginecologia.servicio = ? AND ginecologia.fecha = ?
    GROUP BY catalogoinsumos.nombre"""

    curp = values['-curp-'].upper().strip()
    servicio = values['-servicio-']
    nombre = values['-nombre-'].upper().strip()

    with conn_local:
        consulta = conn_local.execute(query,(curp,servicio,fechahoy))
        for item in consulta:
            rreporte.append(item)
    
    tnom = f'Nombre: {nombre} \nServicio: {servicio}'
    document = Document()
    document.add_heading('Reporte de Material Consumido', 0)
    document.add_heading(tnom, 1)

    for item in rreporte:
        clave, fecha, hora, servicio, nombre, cantidad = item
        f = f'-------------\nClave: {clave} | Fecha: {fecha} | Hora: {hora}\n{nombre}  -  Cantidad: {cantidad}'
        document.add_paragraph(f)

    document.save('./reportes/temporal.docx')

def reporte_pedia():
    """Generar reporte de consumibles pediatria"""

    query = """SELECT pediatria.clave,pediatria.fecha,pediatria.hora,pediatria.servicio,
    catalogoinsumos.nombre, COUNT (catalogoinsumos.nombre) cantidad 
    FROM pediatria,catalogoinsumos 
    WHERE pediatria.paciente_id=?
    AND pediatria.CUPI=? AND pediatria.clave = catalogoinsumos.clave 
    AND ginecologia.servicio = ? AND ginecologia.fecha = ?
    GROUP BY catalogoinsumos.nombre"""

    paciente_id = values['-nombre-'].upper().strip()
    servicio = values['-servicio-']
    CUPI = generar_cupi()


    with conn_local:
        consulta = conn_local.execute(query,(paciente_id,CUPI,servicio,fechahoy))
        for item in consulta:
            rreporte.append(item)
    
    tnom = f'Nombre: {paciente_id} \nServicio: {servicio}'
    document = Document()
    document.add_heading('Reporte de Material Consumido', 0)
    document.add_heading(tnom, 1)

    for item in rreporte:
        clave, fecha, hora, servicio, nombre, cantidad = item
        f = f'-------------\nClave: {clave} | Fecha: {fecha} | Hora: {hora}\n{nombre}  -  Cantidad: {cantidad}'
        document.add_paragraph(f)

    document.save('./reportes/temporal.docx')

def agregar_gine_local():
    valores = (values['-nombre-'].upper().strip(),values['-curp-'].upper().strip(),insumo,fechahoy,tiempo,values['-servicio-'])
    
    query = """INSERT INTO ginecologia (paciente_id,CURP,clave,fecha,hora,servicio) VALUES (?,?,?,?,?,?)"""
    
    for i in range(int(values['-cantidad-'])):
        with conn_local:
            conn_local.execute(query,valores)

def agregar_gine_maria():
    try:
        conn_maria = maria_connect()

        with conn_maria.cursor() as cursor:
            
            valores = (values['-nombre-'].upper().strip(),values['-curp-'].upper().strip(),insumo,fechahoy,tiempo,values['-servicio-'])
        
            query = """INSERT INTO ginecologia (paciente_id,CURP,clave,fecha,hora,servicio) VALUES (%s,%s,%s,%s,%s,%s)"""
    
            for i in range(int(values['-cantidad-'])):
                cursor.execute(query,valores)
            
            conn_maria.commit()

        with conn_local:
            query1 = """SELECT * FROM pediatria ORDER BY ID DESC LIMIT 1"""
            r = conn_local.execute(query1)
            ultimo_registro = r.fetchone()
            
            query2 = """UPDATE pediatria SET synced = 1 WHERE ID = ?"""
            ID = ultimo_registro[0]

            with conn_local:
                conn_local.execute(query2,(ID,))
        
    except:
        pass


def agregar_pedia_local():
    
    CUPI = generar_cupi()
    valores = (values['-nombre-'].upper().strip(),values['-curp-'].upper().strip(),CUPI,insumo,fechahoy,tiempo,values['-servicio-'])
    
    query = """INSERT INTO pediatria (paciente_id,CURP,CUPI,clave,fecha,hora,servicio) VALUES (?,?,?,?,?,?,?)"""
    
    for i in range(int(values['-cantidad-'])):
        with conn_local:
            conn_local.execute(query,valores)


def agregar_pedia_maria():
    
    try:
        conn_maria = maria_connect()

        with conn_maria.cursor() as cursor:
            CUPI = generar_cupi()
            valores = (values['-nombre-'].upper().strip(),values['-curp-'].upper().strip(),CUPI,insumo,fechahoy,tiempo,values['-servicio-']) 

            query = """INSERT INTO pediatria (paciente_id,CURP,CUPI,clave,fecha,hora,servicio) VALUES (%s,%s,%s,%s,%s,%s,%s)"""
    
            for i in range(int(values['-cantidad-'])):
                cursor.execute(query,valores)
            
            conn_maria.commit()

        with conn_local:
            query1 = """SELECT * FROM pediatria ORDER BY ID DESC LIMIT 1"""
            r = conn_local.execute(query1)
            ultimo_registro = r.fetchone()
            
            query2 = """UPDATE pediatria SET synced = 1 WHERE ID = ?"""
            ID = ultimo_registro[0]

            with conn_local:
                conn_local.execute(query2,(ID,))

    except:
        pass


def sincronizar_maria():
    """Funcion para sincronizar registros de BD local con BD maria"""

    #Ambos servicios al mismo tiempo

    try:
        with conn_local:

            cur = conn_local.cursor()

            query = """SELECT paciente_id,CURP,CUPI,clave,fecha,hora,servicio FROM pediatria WHERE synced=0"""

            cur.execute(query)

            lista_pedia = cur.fetchall()

        
            conn_maria = maria_connect()

            with conn_maria.cursor() as cursor:
                query = """INSERT INTO pediatria (paciente_id,CURP,CUPI,clave,fecha,hora,servicio) VALUES (%s,%s,%s,%s,%s,%s,%s)"""

                cursor.executemany(query,lista_pedia)

                conn_maria.commit()


            query2 = """SELECT paciente_id,CURP,clave,fecha,hora,servicio FROM ginecologia WHERE synced=0"""

            cur.execute(query2)

            lista_gine =  cur.fetchall()


            with conn_maria.cursor() as cursor:    

                query = """INSERT INTO ginecologia (paciente_id,CURP,clave,fecha,hora,servicio) VALUES (%s,%s,%s,%s,%s,%s)"""
    
                cursor.executemany(query,lista_gine)

                conn_maria.commit()

            query3 = """SELECT ID FROM pediatria WHERE synced=0"""

            cur.execute(query3)

            sync_pedia =  cur.fetchall()

            query4 = """UPDATE pediatria SET synced = 1 WHERE ID = ?"""

            cur.executemany(query4,sync_pedia)
        
            query5 = """SELECT ID FROM ginecologia WHERE synced=0"""

            cur.execute(query5)

            sync_gine = cur.fetchall()

            query6 = """UPDATE ginecologia SET synced = 1 WHERE ID = ?"""

            cur.executemany(query6,sync_gine)
    
    except:
        pass

    
def busqueda_gine():
    
    curp = values['-curp-'].upper().strip()
    global output_gine
    
    try:
        query = """SELECT ginecologia.clave,ginecologia.fecha, ginecologia.hora,ginecologia.servicio,catalogoinsumos.nombre 
        FROM ginecologia,catalogoinsumos 
        WHERE ginecologia.CURP=%s 
        AND ginecologia.clave = catalogoinsumos.clave 
        ORDER BY ginecologia.ID DESC LIMIT 35"""
        
        conn_maria = maria_connect()

        with conn_maria.cursor() as cursor:
            cursor.execute(query,(curp,))

            lista = cursor.fetchall()
        
            for item in lista:
                output_gine += f'--------------------------------\nClave: {str(item[0])}  |  Fecha: {str(item[1])}  |  Hora: {str(item[2])}  |  Servicio: {str(item[3])}\n\n{str(item[4])}\n'
    
    except:
        pass
    
    
    if output_gine == "":

        query2 = """SELECT ginecologia.clave,ginecologia.fecha, ginecologia.hora,ginecologia.servicio,catalogoinsumos.nombre 
        FROM ginecologia,catalogoinsumos 
        WHERE ginecologia.CURP=?
        AND ginecologia.clave = catalogoinsumos.clave 
        ORDER BY ginecologia.ID DESC LIMIT 35"""

        with conn_local:
            lista = conn_local.execute(query2,(curp,))
        
            for item in lista:
                output_gine += f'--------------------------------\nClave: {str(item[0])}  |  Fecha: {str(item[1])}  |  Hora: {str(item[2])}  |  Servicio: {str(item[3])}\n\n{str(item[4])}\n'

def busqueda_pedia():
    
    nombre = values['-nombre-'].upper().strip()
    CUPI = generar_cupi()
    global output_pedia

    try:

        query = """SELECT pediatria.clave,pediatria.fecha,pediatria.hora,pediatria.servicio,catalogoinsumos.nombre 
        FROM pediatria,catalogoinsumos 
        WHERE pediatria.paciente_id=%s
        AND pediatria.CUPI=%s
        AND pediatria.clave = catalogoinsumos.clave 
        ORDER BY pediatria.ID DESC LIMIT 35"""

        conn_maria = maria_connect()

        with conn_maria.cursor() as cursor:
            cursor.execute(query,(nombre,CUPI))

            lista = cursor.fetchall()
    
            for item in lista:
                output_pedia += f'--------------------------------\nClave: {str(item[0])}  |  Fecha: {str(item[1])}  |  Hora: {str(item[2])}  | Servicio: {str(item[3])}\n\n{str(item[4])}\n'

    except:
        pass

    if output_pedia == "":

        query2 = """SELECT pediatria.clave,pediatria.fecha,pediatria.hora,pediatria.servicio,catalogoinsumos.nombre 
            FROM pediatria,catalogoinsumos 
            WHERE pediatria.paciente_id=?
            AND pediatria.CUPI=? 
            AND pediatria.clave = catalogoinsumos.clave 
            ORDER BY pediatria.ID DESC LIMIT 35"""

        with conn_local:
            lista = conn_local.execute(query2,(nombre,CUPI))
    
            for item in lista:
                output_pedia += f'--------------------------------\nClave: {str(item[0])}  |  Fecha: {str(item[1])}  |  Hora: {str(item[2])}  | Servicio: {str(item[3])}\n\n{str(item[4])}\n'
            
def busqueda_catalogoinsumos():
    query="""SELECT clave,nombre FROM catalogoinsumos WHERE nombre LIKE ?"""

    palabraclave = values['-busnom-'].lower().strip().replace('de','').replace('con','').replace('para','').split()
    
    for palabra in palabraclave:
        if palabra.isdigit():
            pass
        else:
            input_usuario = '%' + palabra  + '%'
            with conn_local:
                lista = conn_local.execute(query,(input_usuario,))
    
                for item in lista:
                    if item[1] in alista:
                        pass
                    else:
                        alista.append(item[1])

def generar_cupi():
    """Funcion para generar el CUPI (Clave Unica Pediatrica de Identificacion)"""
    apellidos = values['-nombre-'].upper().strip().replace('RN','').split()
    fn = values['-fn-'].replace('/','')
    if values['-masc-']:
        sexo = "H"
    elif values['-fem-']:
        sexo = "M"
    CUPI = f'{apellidos[0][0:2]}{apellidos[1][0:2]}{fn[2:]}{sexo}'
    return CUPI


#Interfaz Grafica - PySimpleGUI
sg.theme('Teal Mono')  # Colores!

#Elementos en ventana

menuprincipal = [
    ['Inventario', ['Capturar Existencias', 'Capturar Salidas']],
    ['Ayuda', ['Instrucciones','Acerca de...' ]]
]

Info_Paciente = [
    [sg.Text('Nombre del paciente'), sg.Input(size=(40,2), key='-nombre-'),sg.Text('CURP'), sg.Input(size=(20,2), key='-curp-')],
    [ sg.Text('Servicio'), 
     sg.Combo(['Urgencias GyO','Labor / Toco (GyO)','Quirofano GyO','Alojamiento Conjunto GyO','Modulo Mater','Urgencias Pediatria',
    'Labor / Toco (Pediatria)', 'Quirofano Pediatria','UCIN','Cunero Patologico','Cuidados Intermedios',
    'Alojamiento Conjunto Pediatria'], size=(25,12), readonly=True, key='-servicio-'), sg.CalendarButton('Fecha de Nacimiento', target=(1,3), auto_size_button=True, format='%Y/%m/%d'), 
     sg.Input(size=(13,2), key='-fn-'), sg.Radio('Masculino', "RADIO1",key='-masc-'), sg.Radio('Femenino', "RADIO1",default=True, key='-fem-')]
]

AreaCaptura = [
    [sg.Text('Busqueda por palabras clave')],
    [sg.Input(size=(55,2), key='-busnom-'), sg.Button('Buscar insumo'), sg. Button('Borrar Seleccion')],
    [sg.Listbox(values='',size=(100, 8), key='-linsumo-')],
    [sg.Text('Busqueda manual')],
    [sg.Combo(mcombo, size=(100,6), key='-minsumo-')],
    [sg.Text('Cantidad de material'), sg.Input(default_text='1', size=(5,2), key='-cantidad-'), sg.Button('Agregar Registro')]
]

AreaRegistros = [
    [sg.Text('Ultimos movimientos de material registrados'), sg.Text('\t\t\t'), sg.Button('Buscar Registros'), sg.Button('Imprimir Reporte')],
    [sg.Multiline(size=(100,12),disabled=True,key='-output-')]
]

# Acomodo de elementos de ventana
layout = [
    [sg.Menu(menuprincipal, tearoff=True)],
    [sg.Frame('Identificacion del paciente', Info_Paciente)],
    [sg.Frame('Busqueda y captura de material hospitalario', AreaCaptura)],
    [sg.Frame('Consulta de movimientos', AreaRegistros)]
]

#Ventana
window = sg.Window('HMI - Captura de datos', layout)

#Sincronizar registros pendientes despues de iniciar la ventana
with concurrent.futures.ThreadPoolExecutor(max_workers=10) as executor:
    executor.submit(sincronizar_maria())


#Parte funcional de la aplicacion

while True:  # Mantenerla persistente

    event, values = window.read()

    if event == 'Capturar Existencias':
        print(event,values)

    if event == sg.WIN_CLOSED:
        break
    
    #Buscar insumos
    elif event == 'Buscar insumo':
        if values['-busnom-'] != '':
            busqueda_catalogoinsumos()
            window['-linsumo-'].update(alista)
            alista = []
        else:
            sg.popup('¿CAW?','No olvides escribir en la casilla de busqueda\n')

    #borrar seleccion
    elif event == 'Borrar Seleccion':
        alista = []
        window['-linsumo-'].update(alista)
        
    #Buscar registros
    elif event == 'Buscar Registros':
        if values['-servicio-'] == 'Urgencias GyO' or values['-servicio-'] == 'Labor / Toco (GyO)' or values['-servicio-'] == 'Quirofano GyO' or values['-servicio-'] == 'Alojamiento Conjunto GyO' or values['-servicio-'] == 'Modulo Mater':
            if values['-nombre-'] != '' and values['-curp-'] != '':
                busqueda_gine()
                
                if output_gine == "":
                    sg.popup('Baia Baia','El paciente no existe en la base de datos\n\nIntenta de nuevo\n\nTip: Verifica el nombre, curp o servicio\n')
                else:
                    window['-output-'].update(output_gine)
                    output_gine = ""
            else:
                    sg.popup('¿CAW?','No olvides los datos del paciente! (Nombre y CURP)\n')

        elif values['-servicio-'] == 'Urgencias Pediatria' or values['-servicio-'] == 'Labor / Toco (Pediatria)' or values['-servicio-'] == 'Quirofano Pediatria' or values['-servicio-'] == 'UCIN' or values['-servicio-'] == 'Cunero Patologico' or values['-servicio-'] =='Cuidados Intermedios' or values['-servicio-'] == 'Alojamiento Conjunto Pediatria':
            if values['-nombre-'] != '' and len(values['-fn-']) == 10:
                busqueda_pedia()
                
                if output_pedia == "":
                    sg.popup('Baia Baia','El paciente no existe en la base de datos\n\nIntenta de nuevo\n\nTip: Verifica que el nombre, servicio, fecha de nacimiento y sexo sean correctos\n')
                else:
                    window['-output-'].update(output_pedia)
                    output_pedia = ""
            else:
                sg.popup('¿CAW?','No olvides escribir el nombre y la fecha de nacimiento (CURP es opcional)!\n')
        else:
            sg.popup('¿CAW?','No olvides seleccionar un servicio\n')       
    
    #Imprimir Reporte
    elif event == 'Imprimir Reporte':
        if values['-servicio-'] == 'Urgencias GyO' or values['-servicio-'] == 'Labor / Toco (GyO)' or values['-servicio-'] == 'Quirofano GyO' or values['-servicio-'] == 'Alojamiento Conjunto GyO' or values['-servicio-'] == 'Modulo Mater':
            if values['-nombre-'] != '' and values['-curp-'] != '':
                reporte_gine()
                rreporte = []
            else:
                sg.popup('¿CAW?','No olvides los datos del paciente! (Nombre y CURP)\n')
        
        elif values['-servicio-'] == 'Urgencias Pediatria' or values['-servicio-'] == 'Labor / Toco (Pediatria)' or values['-servicio-'] == 'Quirofano Pediatria' or values['-servicio-'] == 'UCIN' or values['-servicio-'] == 'Cunero Patologico' or values['-servicio-'] =='Cuidados Intermedios' or values['-servicio-'] == 'Alojamiento Conjunto Pediatria':
            if values['-nombre-'] != '' and len(values['-fn-']) == 10:
                reporte_pedia()
                rreporte = []            
            else:
                sg.popup('¿CAW?','No olvides los datos del paciente! (Nombre y Fecha Nacimiento)\n')
        else:
            sg.popup('¿CAW?','No olvides seleccionar un servicio\n')

    #Enviar Datos
    elif event == 'Agregar Registro':
        if values['-servicio-'] == 'Urgencias GyO' or values['-servicio-'] == 'Labor / Toco (GyO)' or values['-servicio-'] == 'Quirofano GyO' or values['-servicio-'] == 'Alojamiento Conjunto GyO' or values['-servicio-'] == 'Modulo Mater':
            if values['-nombre-'] != '' and values['-curp-'] != '' and values['-cantidad-'].isdigit() and int(values['-cantidad-']) < 20:
                if values['-linsumo-'] != []:
                    seleccion = values['-linsumo-']
                    insumo = catalogo[seleccion[0]]
                    agregar_gine_local()
                    agregar_gine_maria()
                    busqueda_gine()
                    window['-output-'].update(output_gine)
                    output_gine = ""
                    

                elif values['-minsumo-'] != '':
                    try:
                        llave = values['-minsumo-']
                        insumo = catalogo[llave]
                        agregar_gine_local()
                        agregar_gine_maria()
                        busqueda_gine()
                        window['-output-'].update(output_gine)
                        output_gine = ""
                    except:
                        sg.popup('¿CAW?','Parece que la clave esta incompleta, selecciona de nuevo dentro de la lista\n')
           
                else:
                    sg.popup('¿CAW?','No olvides seleccionar una opcion de insumo\n')
            else:
                sg.popup('¿CAW? Recuerda...','-Necesito nombre y curp! No puedo registrarlo asi!\n\n-Tampoco debes usar una letra como cantidad!\n\nNo es algebra!\n\n-No puedes registrar mas de 20 articulos a la vez!\n\nMuere la computadora!\n\n')
       
        elif values['-servicio-'] == 'Urgencias Pediatria' or values['-servicio-'] == 'Labor / Toco (Pediatria)' or values['-servicio-'] == 'Quirofano Pediatria' or values['-servicio-'] == 'UCIN' or values['-servicio-'] == 'Cunero Patologico' or values['-servicio-'] =='Cuidados Intermedios' or values['-servicio-'] == 'Alojamiento Conjunto Pediatria':
            if values['-nombre-'] != '' and len(values['-fn-']) == 10 and values['-cantidad-'].isdigit() and int(values['-cantidad-']) < 20:
                if values['-linsumo-'] != []:
                    seleccion = values['-linsumo-']
                    insumo = catalogo[seleccion[0]]
                    agregar_pedia_local()
                    agregar_pedia_maria()
                    busqueda_pedia()
                    window['-output-'].update(output_pedia)
                    output_pedia = ""

                elif values['-minsumo-'] != '':
                    try:
                        llave = values['-minsumo-']
                        insumo = catalogo[llave]
                        agregar_pedia_local()
                        agregar_pedia_maria()
                        busqueda_pedia()
                        window['-output-'].update(output_pedia)
                        output_pedia = ""
                    except:
                        sg.popup('¿CAW?','Parece que el articulo esta incompleto, selecciona de nuevo dentro de la lista\n')
                else:
                    sg.popup('¿CAW?','No olvides seleccionar una opcion de insumo\n')
            else:
                sg.popup('¿CAW? Recuerda...','-Necesito un nombre y una fecha de nacimiento! No puedo registrarlo asi!\n\n(El CURP es opcional)\n\n-Tampoco debes usar una letra como cantidad!\n\nNo es algebra!\n\n-No puedes registrar mas de 20 articulos a la vez!\n\nMuere la computadora!\n\n')   
        else:
            sg.popup('¿CAW?','No haz elegido servicio! no puedo trabajar asi\n')

window.close()